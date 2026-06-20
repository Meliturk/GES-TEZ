from pathlib import Path
import re

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT = Path(r"C:\Users\TURK PC\Desktop\GES_\GES_AI")
DOCS = PROJECT / "docs"
OUT = DOCS / "GES_TEZ_TASLAK.docx"


def set_cell_text(cell, text, bold=False, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)


def set_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for style_name, size in [
        ("Heading 1", 14),
        ("Heading 2", 13),
        ("Heading 3", 12),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    if "EK Kod" not in doc.styles:
        code_style = doc.styles.add_style("EK Kod", 1)
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(8)
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")


def format_paragraph(p, alignment=None, line_spacing=1.5, space_after=6):
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()

    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_cover(doc):
    lines = [
        ("BURSA ULUDAĞ ÜNİVERSİTESİ", 14, True),
        ("İNEGÖL İŞLETME FAKÜLTESİ", 14, True),
        ("YÖNETİM BİLİŞİM SİSTEMLERİ BÖLÜMÜ", 14, True),
        ("", 12, False),
        ("", 12, False),
        ("GÜNEŞ ENERJİ SANTRALLERİNDE BÜYÜK VERİ VE YAPAY ZEKÂ TABANLI ÜRETİM TAHMİNİ, ANOMALİ TESPİTİ VE AKILLI ENERJİ YÖNETİMİ", 14, True),
        ("", 12, False),
        ("", 12, False),
        ("LİSANS TEZİ", 14, True),
        ("", 12, False),
        ("", 12, False),
        ("Hazırlayan", 12, False),
        ("Melih Türk", 12, True),
        ("132230091", 12, False),
        ("", 12, False),
        ("Danışman", 12, False),
        ("Melih Engin", 12, True),
        ("", 12, False),
        ("", 12, False),
        ("Bursa", 12, False),
        ("2026", 12, False),
    ]

    for text, size, bold in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        format_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, 1.5, 8)

    doc.add_page_break()


def classify_line(line):
    s = line.strip()

    if not s:
        return "blank"

    if s.startswith("=") or s.startswith("---"):
        return "separator"

    if re.match(r"^\d+\.\s+[A-ZÇĞİÖŞÜ0-9]", s):
        return "h1"

    if re.match(r"^\d+\.\d+\.\s+[A-ZÇĞİÖŞÜ0-9]", s):
        return "h2"

    if re.match(r"^\d+\.\d+\.\d+\.\s+[A-ZÇĞİÖŞÜ0-9]", s):
        return "h3"

    upper_titles = [
        "ÖZET",
        "ABSTRACT",
        "KAYNAKÇA",
        "EKLER",
        "GİRİŞ",
        "BULGULAR",
    ]

    if s in upper_titles:
        return "h1"

    if s.startswith("Şekil "):
        return "caption"

    if s.startswith("- "):
        return "bullet"

    return "normal"


def add_text_file(doc, file_path, appendix=False):
    file_path = Path(file_path)

    if not file_path.exists():
        p = doc.add_paragraph(f"[DOSYA BULUNAMADI: {file_path}]")
        format_paragraph(p)
        return

    text = file_path.read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()

    for line in lines:
        line_type = classify_line(line)

        if line_type == "blank":
            doc.add_paragraph("")
            continue

        if line_type == "separator":
            continue

        if appendix:
            if line.strip().startswith("EK-") or line.strip() in ["EKLER"]:
                p = doc.add_paragraph(line.strip(), style="Heading 2")
                format_paragraph(p, None, 1.15, 8)
            else:
                p = doc.add_paragraph(line, style="EK Kod")
                format_paragraph(p, None, 1.0, 0)
            continue

        if line_type == "h1":
            p = doc.add_paragraph(line.strip(), style="Heading 1")
            format_paragraph(p, None, 1.15, 12)

        elif line_type == "h2":
            p = doc.add_paragraph(line.strip(), style="Heading 2")
            format_paragraph(p, None, 1.15, 8)

        elif line_type == "h3":
            p = doc.add_paragraph(line.strip(), style="Heading 3")
            format_paragraph(p, None, 1.15, 6)

        elif line_type == "caption":
            p = doc.add_paragraph(line.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].italic = True
            format_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, 1.15, 8)

        elif line_type == "bullet":
            p = doc.add_paragraph(line.strip()[2:], style="List Bullet")
            format_paragraph(p, None, 1.5, 4)

        else:
            p = doc.add_paragraph(line.strip())
            format_paragraph(p, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.5, 6)


def add_image_with_caption(doc, img_path, caption, width_cm=15.5):
    img_path = Path(img_path)

    if not img_path.exists():
        p = doc.add_paragraph(f"[GÖRSEL BULUNAMADI: {img_path}]")
        format_paragraph(p)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
    format_paragraph(cap, WD_ALIGN_PARAGRAPH.CENTER, 1.15, 10)


def add_figures_section(doc):
    doc.add_page_break()
    p = doc.add_paragraph("5.14. Bulgulara Ait Grafikler ve Dashboard Görselleri", style="Heading 2")
    format_paragraph(p, None, 1.15, 12)

    figures = [
        ("final_outputs/anomaly_by_source.png", "Şekil 5.1. Plant 1 kaynak bazlı anomali dağılımı"),
        ("final_outputs/anomaly_by_day.png", "Şekil 5.2. Plant 1 günlük anomali dağılımı"),
        ("final_outputs/actual_vs_expected_top_source.png", "Şekil 5.3. Plant 1 beklenen ve gerçekleşen AC güç karşılaştırması"),
        ("final_outputs/plant_2_anomaly_by_source.png", "Şekil 5.4. Plant 2 kaynak bazlı anomali dağılımı"),
        ("final_outputs/plant_2_anomaly_by_day.png", "Şekil 5.5. Plant 2 günlük anomali dağılımı"),
        ("final_outputs/plant_2_actual_vs_expected.png", "Şekil 5.6. Plant 2 beklenen ve gerçekleşen AC güç karşılaştırması"),
        ("final_outputs/plant_anomaly_rate_comparison.png", "Şekil 5.7. Plant 1 ve Plant 2 anomali oranı karşılaştırması"),
        ("final_outputs/plant_model_r2_comparison.png", "Şekil 5.8. Plant 1 ve Plant 2 model R² karşılaştırması"),
        ("final_outputs/real_consumption_production_vs_consumption_sample.png", "Şekil 5.9. Gerçek tüketim destekli üretim ve tüketim karşılaştırması"),
        ("final_outputs/real_consumption_decision_distribution.png", "Şekil 5.10. Gerçek tüketim destekli enerji yönetimi karar dağılımı"),
        ("final_outputs/real_consumption_battery_soc_simulation.png", "Şekil 5.11. Akü SOC simülasyonu"),
        ("final_outputs/real_consumption_grid_sell_energy_simulation.png", "Şekil 5.12. Şebekeye satışa uygun enerji simülasyonu"),
        ("docs/tez_gorselleri/buyuk_veri_laboratuvari/dashboard_01_ozet_kartlar_ve_uyarilar.jpeg", "Şekil 5.13. GES Kafka canlı dashboard özet kartları ve uyarı alanı"),
        ("docs/tez_gorselleri/buyuk_veri_laboratuvari/dashboard_02_grafikler.jpeg", "Şekil 5.14. Kafka dashboard üzerinde son 1 saate ait sıcaklık, nem, ışık şiddeti, akım ve raw ADC grafikleri"),
        ("docs/tez_gorselleri/buyuk_veri_laboratuvari/dashboard_03_topic_son_mesaj_tablolari.jpeg", "Şekil 5.15. Kafka topiclerinden gelen ENV, PANEL ve COMBINED son mesaj tabloları"),
        ("docs/tez_gorselleri/buyuk_veri_laboratuvari/dashboard_04_anomali_uyari_ornegi.jpeg", "Şekil 5.16. Dashboard üzerinde ADC saturasyon ve sensör verisi uyarı örnekleri"),
    ]

    for rel_path, caption in figures:
        add_image_with_caption(doc, PROJECT / rel_path, caption)
        doc.add_paragraph("")


def build_docx():
    doc = Document()
    set_document_styles(doc)

    for section in doc.sections:
        add_page_number(section)

    add_cover(doc)

    add_text_file(doc, DOCS / "TEZ_OZET.txt")
    doc.add_page_break()

    add_text_file(doc, DOCS / "TEZ_ABSTRACT.txt")
    doc.add_page_break()

    add_text_file(doc, DOCS / "TEZ_1_GIRIS.txt")
    doc.add_page_break()

    add_text_file(doc, DOCS / "TEZ_2_TEORIK_ARKA_PLAN.txt")
    doc.add_page_break()

    add_text_file(doc, DOCS / "TEZ_3_LITERATUR_TARAMASI.txt")
    doc.add_page_break()

    add_text_file(doc, DOCS / "TEZ_4_YONTEM.txt")
    doc.add_page_break()

    add_text_file(doc, DOCS / "TEZ_5_BULGULAR.txt")
    add_figures_section(doc)
    doc.add_page_break()

    add_text_file(doc, DOCS / "TEZ_6_TARTISMA_VE_SONUC.txt")
    doc.add_page_break()

    add_text_file(doc, DOCS / "TEZ_KAYNAKCA.txt")
    doc.add_page_break()

    add_text_file(doc, DOCS / "TEZ_EKLER.txt", appendix=True)

    doc.save(OUT)
    print("DOCX oluşturuldu:", OUT)


if __name__ == "__main__":
    build_docx()
